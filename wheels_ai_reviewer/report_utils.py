"""
report_utils.py — Final Report generation (DOCX + PDF)
=========================================================
Builds the downloadable Final Report shown on the Report step of the wizard:
    - Overall dashboard (risk score, pass/flag/fail counts)
    - Business impact summary (from the Business Impact step)
    - Reviewer decisions (Accept/Reject) with captured reasons
    - Full findings list, grouped by checklist category

Two independent renderers are provided so the report is a straight
translation of the same `report_data` dict into either format:
    build_docx_report(report_data) -> bytes   (python-docx)
    build_pdf_report(report_data)  -> bytes   (reportlab)

Install:
    pip install python-docx reportlab

`report_data` shape (see report_page() in app.py):
{
    "contract_name": str,
    "generated_at": str,
    "score": int,
    "fail_n": int, "flag_n": int, "pass_n": int,
    "findings": list[Finding],           # from msa_reviewer_agent
    "decisions": dict[str, dict],        # decision_id -> {category, checklist_item,
                                          #   status, severity, clause_reference, issue,
                                          #   recommendation, decision, reason}
    "business_impact": dict,             # from roi_page(), may be {}
}
"""

from io import BytesIO
from xml.sax.saxutils import escape as _esc

# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

STATUS_LABEL = {"pass": "PASS", "flag": "FLAG", "fail": "FAIL"}


def _style_table(table):
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        # Style not available in the base template — leave default borders.
        pass


def build_docx_report(report_data: dict) -> bytes:

    findings = report_data.get("findings") or []
    decisions = report_data.get("decisions") or {}
    bi = report_data.get("business_impact") or {}

    doc = Document()

    title = doc.add_heading("MSA Review — Final Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Contract: ").bold = True
    meta.add_run(str(report_data.get("contract_name", "Uploaded Contract")))
    meta2 = doc.add_paragraph()
    meta2.add_run("Generated: ").bold = True
    meta2.add_run(str(report_data.get("generated_at", "")))

    # ---------------- Overall Dashboard ----------------
    doc.add_heading("Overall Dashboard", level=1)

    dash = doc.add_table(rows=2, cols=4)
    _style_table(dash)
    labels = ["Risk Score", "Failures", "Flags", "Passed"]
    values = [
        report_data.get("score", 0),
        report_data.get("fail_n", 0),
        report_data.get("flag_n", 0),
        report_data.get("pass_n", 0),
    ]
    for i, label in enumerate(labels):
        dash.rows[0].cells[i].text = label
        dash.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        dash.rows[1].cells[i].text = str(values[i])

    # ---------------- Business Impact ----------------
    if bi:

        doc.add_heading("Business Impact", level=1)

        bi_rows = [
            ("Reviewer Hours Saved (this contract)", f"{bi.get('hours_saved_per_contract', 0):.1f} hrs"),
            ("Review Cost Saved (this contract)", f"${bi.get('cost_saved_per_contract', 0):,.0f}"),
            ("Risk Exposure Flagged", f"${bi.get('risk_exposure_avoided', 0):,.0f}"),
            ("Total Business Impact (this contract)", f"${bi.get('total_impact', 0):,.0f}"),
            ("Annual Review Cost Saved (at current volume)", f"${bi.get('annual_cost_saved', 0):,.0f}"),
        ]

        bi_table = doc.add_table(rows=0, cols=2)
        _style_table(bi_table)
        for label, val in bi_rows:
            cells = bi_table.add_row().cells
            cells[0].text = label
            cells[1].text = val

        doc.add_paragraph(
            "These are directional estimates based on the assumptions entered "
            "in the Business Impact step, not audited figures.",
        ).runs[0].italic = True

    # ---------------- Reviewer Decisions ----------------
    doc.add_heading("Reviewer Decisions", level=1)

    reviewed = [d for d in decisions.values() if d.get("decision") in ("Accept", "Reject")]

    if not reviewed:
        doc.add_paragraph("No items were explicitly accepted or rejected during review.")
    else:
        for d in reviewed:
            p = doc.add_paragraph()
            run = p.add_run(f"{d['decision'].upper()} — {d['checklist_item']}")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x7F, 0x37) if d["decision"] == "Accept" else RGBColor(0xC0, 0x2B, 0x2B)

            doc.add_paragraph(
                f"Category: {d['category'].replace('_', ' ').title()}    "
                f"Severity: {d['severity'].title()}    "
                f"Clause: {d['clause_reference']}"
            )
            doc.add_paragraph(f"Issue: {d['issue']}")
            doc.add_paragraph(f"Recommendation: {d['recommendation']}")

            if d.get("reason"):
                reason_p = doc.add_paragraph()
                reason_run = reason_p.add_run(f"Reason: {d['reason']}")
                reason_run.italic = True

            doc.add_paragraph("")

    # ---------------- All Findings ----------------
    doc.add_heading("All Findings", level=1)

    categories = sorted(set(f.category for f in findings))

    if not categories:
        doc.add_paragraph("No findings recorded.")

    for cat in categories:
        doc.add_heading(cat.replace("_", " ").title(), level=2)
        for f in [x for x in findings if x.category == cat]:
            p = doc.add_paragraph()
            run = p.add_run(f"[{STATUS_LABEL.get(f.status, f.status.upper())}] {f.checklist_item}")
            run.bold = True
            doc.add_paragraph(f"Severity: {f.severity.title()}    Clause: {f.clause_reference}")
            doc.add_paragraph(f.issue)
            doc.add_paragraph(f"Recommendation: {f.recommendation}")
            doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)

BRAND_BLUE = colors.HexColor("#004990")
LIGHT_BG = colors.HexColor("#F4F7FB")
ACCEPT_COLOR = "#1a7f37"
REJECT_COLOR = "#c02b2b"


def build_pdf_report(report_data: dict) -> bytes:

    findings = report_data.get("findings") or []
    decisions = report_data.get("decisions") or {}
    bi = report_data.get("business_impact") or {}

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=48, bottomMargin=48)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("MSA Review — Final Report", styles["Title"]))
    story.append(Paragraph(f"Contract: {_esc(str(report_data.get('contract_name', 'Uploaded Contract')))}", styles["Normal"]))
    story.append(Paragraph(f"Generated: {_esc(str(report_data.get('generated_at', '')))}", styles["Normal"]))
    story.append(Spacer(1, 16))

    # ---------------- Overall Dashboard ----------------
    story.append(Paragraph("Overall Dashboard", styles["Heading1"]))

    dash_table = Table(
        [
            ["Risk Score", "Failures", "Flags", "Passed"],
            [
                str(report_data.get("score", 0)),
                str(report_data.get("fail_n", 0)),
                str(report_data.get("flag_n", 0)),
                str(report_data.get("pass_n", 0)),
            ],
        ],
        hAlign="LEFT",
    )
    dash_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BRAND_BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(dash_table)
    story.append(Spacer(1, 16))

    # ---------------- Business Impact ----------------
    if bi:

        story.append(Paragraph("Business Impact", styles["Heading1"]))

        bi_rows = [
            ["Reviewer Hours Saved (this contract)", f"{bi.get('hours_saved_per_contract', 0):.1f} hrs"],
            ["Review Cost Saved (this contract)", f"${bi.get('cost_saved_per_contract', 0):,.0f}"],
            ["Risk Exposure Flagged", f"${bi.get('risk_exposure_avoided', 0):,.0f}"],
            ["Total Business Impact (this contract)", f"${bi.get('total_impact', 0):,.0f}"],
            ["Annual Review Cost Saved (at current volume)", f"${bi.get('annual_cost_saved', 0):,.0f}"],
        ]
        bi_table = Table(bi_rows, hAlign="LEFT", colWidths=[280, 160])
        bi_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), LIGHT_BG),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(bi_table)
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            "<i>These are directional estimates based on the assumptions entered "
            "in the Business Impact step, not audited figures.</i>",
            styles["Normal"],
        ))
        story.append(Spacer(1, 16))

    # ---------------- Reviewer Decisions ----------------
    story.append(Paragraph("Reviewer Decisions", styles["Heading1"]))

    reviewed = [d for d in decisions.values() if d.get("decision") in ("Accept", "Reject")]

    if not reviewed:
        story.append(Paragraph("No items were explicitly accepted or rejected during review.", styles["Normal"]))
    else:
        for d in reviewed:
            color_hex = ACCEPT_COLOR if d["decision"] == "Accept" else REJECT_COLOR
            story.append(Paragraph(
                f"<b><font color='{color_hex}'>{_esc(d['decision'].upper())}</font> — "
                f"{_esc(d['checklist_item'])}</b>",
                styles["Normal"],
            ))
            story.append(Paragraph(
                f"Category: {_esc(d['category'].replace('_', ' ').title())} | "
                f"Severity: {_esc(d['severity'].title())} | "
                f"Clause: {_esc(str(d['clause_reference']))}",
                styles["Normal"],
            ))
            story.append(Paragraph(f"Issue: {_esc(d['issue'])}", styles["Normal"]))
            story.append(Paragraph(f"Recommendation: {_esc(d['recommendation'])}", styles["Normal"]))
            if d.get("reason"):
                story.append(Paragraph(f"<i>Reason: {_esc(d['reason'])}</i>", styles["Normal"]))
            story.append(Spacer(1, 10))

    # ---------------- All Findings ----------------
    story.append(PageBreak())
    story.append(Paragraph("All Findings", styles["Heading1"]))

    categories = sorted(set(f.category for f in findings))

    if not categories:
        story.append(Paragraph("No findings recorded.", styles["Normal"]))

    for cat in categories:
        story.append(Paragraph(_esc(cat.replace("_", " ").title()), styles["Heading2"]))
        for f in [x for x in findings if x.category == cat]:
            story.append(Paragraph(
                f"[{_esc(f.status.upper())}] {_esc(f.checklist_item)}", styles["Normal"]
            ))
            story.append(Paragraph(
                f"Severity: {_esc(f.severity.title())} | Clause: {_esc(str(f.clause_reference))}",
                styles["Normal"],
            ))
            story.append(Paragraph(_esc(f.issue), styles["Normal"]))
            story.append(Paragraph(f"Recommendation: {_esc(f.recommendation)}", styles["Normal"]))
            story.append(Spacer(1, 8))

    doc.build(story)
    return buf.getvalue()
